# -- coding: utf-8 --
import time
start_time = time.time()
import pandas as pd
import os
from docx import Document
from tqdm import tqdm

folder_path = "奖状"
if not os.path.isdir(folder_path):
    os.mkdir(folder_path)

# 读取Excel文件
df = pd.read_excel('3完整表.xlsx', sheet_name='Sheet1')

for i in tqdm(range(len(df["姓名"]))):

    doc = Document("demo.docx")

    doc.add_paragraph('奖状', style=doc.styles['Title'])
    doc.add_paragraph(f'尊敬的{df["姓名"][i]}同学：', style=doc.styles['论文正文'])
    doc.add_paragraph(f'你在本次考试中表现优异，获得了{df["成绩"][i]}分的好成绩。在此，学校颁发给你此奖状一份，以资鼓励！', style=doc.styles['论文正文'])
    doc.add_paragraph(f'今天你勇夺{df["奖项"][i]}，不仅是对自己优异成绩的认可，更是对自己勤奋努力的肯定。希望你在以后的学习和生活中，能够保持这份努力和自信，为自己的梦想不懈地追求。祝贺你！', style=doc.styles['论文正文'])

    doc.save(f'奖状\\{df["姓名"][i]}.docx')

end_time = time.time()
total_time = end_time - start_time
print("\033[31m程序运行时间为：{:.2f}秒\033[0m".format(total_time))